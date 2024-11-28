import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown
from bs4 import BeautifulSoup
import logging
from html2docx import html2docx

logger = logging.getLogger(__name__)

class WordGenerator:
    def __init__(self):
        self.document = Document()
        self._setup_document()

    def _setup_document(self):
        """Setup document styles and formatting"""
        # Set up styles
        style = self.document.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        # Add title
        title = self.document.add_heading('Generated Cold Emails', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add generation timestamp
        timestamp = self.document.add_paragraph()
        timestamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        timestamp.add_run(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        
        # Add separator
        self.document.add_paragraph('=' * 50)

    def _markdown_to_docx(self, markdown_text: str, paragraph) -> None:
        """Convert markdown text to docx formatting"""
        # Convert markdown to HTML
        html = markdown.markdown(markdown_text)
        soup = BeautifulSoup(html, 'html.parser')

        # Process each element
        for element in soup.children:
            if element.name == 'p':
                for content in element.children:
                    if content.name == 'a':
                        # Handle links
                        run = paragraph.add_run(content.text)
                        run.underline = True
                        run.font.color.rgb = None  # Set to blue
                    elif content.name == 'strong':
                        # Handle bold text
                        run = paragraph.add_run(content.text)
                        run.bold = True
                    elif content.name == 'em':
                        # Handle italic text
                        run = paragraph.add_run(content.text)
                        run.italic = True
                    else:
                        # Handle plain text
                        paragraph.add_run(str(content))
            elif element.name == 'hr':
                # Handle horizontal rules
                paragraph.add_run('_' * 50)

    def add_email(self, company_name: str, email_content: str) -> None:
        """Add an email to the document with proper formatting"""
        try:
            # Add company header
            self.document.add_heading(f'Email for: {company_name}', level=1)
            
            # Add email content
            email_paragraph = self.document.add_paragraph()
            self._markdown_to_docx(email_content, email_paragraph)
            
            # Add separator
            self.document.add_paragraph('=' * 50)
            
        except Exception as e:
            logger.error(f"Error adding email for {company_name}: {str(e)}")

    def generate_all_emails_doc(self, profiles, output_path='generated_emails.docx'):
        """Generate a Word document containing all emails"""
        try:
            if not profiles:
                logger.warning("No profiles provided to generate document")
                return False

            for company_name, email_text in profiles:
                # Add company name as heading
                self.document.add_heading(f'Email for {company_name}', level=1)
                
                # Add email content
                email_para = self.document.add_paragraph()
                self._markdown_to_docx(email_text, email_para)
                
                # Add separator
                self.document.add_paragraph('=' * 50)

            # Save the document
            self.document.save(output_path)
            logger.info(f"Successfully generated Word document at {output_path}")
            return True

        except Exception as e:
            logger.error(f"Error generating Word document: {str(e)}")
            return False

    def generate_from_html(self, html_path, output_path='generated_emails.docx'):
        """Generate Word document from HTML file"""
        try:
            # Convert HTML to DOCX
            html2docx(html_path, output_path)
            logger.info(f"Successfully generated Word document at {output_path}")
            return True
        except Exception as e:
            logger.error(f"Error converting HTML to Word: {str(e)}")
            return False

    def save(self, output_path: str) -> str:
        """Save the document and return the path"""
        try:
            # Ensure the directory exists
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            # Save the document
            self.document.save(output_path)
            return output_path
            
        except Exception as e:
            logger.error(f"Error saving document: {str(e)}")
            return None
